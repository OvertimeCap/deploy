from docx import Document
from docx.shared import Pt
import io

def generate_contract_docx(template_path, output_path_or_buffer, context):
    """
    Gera um contrato DOCX a partir de um modelo, substituindo as variáveis.
    :param template_path: Caminho para o arquivo de modelo .docx.
    :param output_path_or_buffer: Caminho para salvar o arquivo gerado ou um buffer (io.BytesIO).
    :param context: Dicionário com as variáveis e seus valores para substituição.
                    Ex: {"{{CLIENTE}}": "Nome do Cliente", "{{CPF}}": "123.456.789-00"}
    """
    try:
        document = Document(template_path)

        # Substituição em parágrafos
        for paragraph in document.paragraphs:
            for key, value in context.items():
                if key in paragraph.text:
                    # Mantém a formatação original o máximo possível, substituindo apenas o texto
                    # Esta é uma abordagem simples. Para formatações complexas dentro de um placeholder,
                    # pode ser necessário dividir o run e aplicar formatações específicas.
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, str(value) if value is not None else "")
                            inline[i].text = text
        
        # Substituição em tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in context.items():
                            if key in paragraph.text:
                                inline = paragraph.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, str(value) if value is not None else "")
                                        inline[i].text = text
        
        # Salva o documento no caminho ou buffer especificado
        document.save(output_path_or_buffer)
        return True
    except Exception as e:
        print(f"Erro ao gerar contrato DOCX: {e}")
        return False

# Exemplo de uso (para teste local, não será executado diretamente pelo agente aqui)
if __name__ == "__main__":
    mock_context = {
        "{{CLIENTE}}": "Fulano de Tal da Silva",
        "{{CPF}}": "111.222.333-44",
        "{{RG}}": "12.345.678-9",
        "{{ENDERECO}}": "Rua das Palmeiras, 123, Bairro Feliz, Cidade Alegre - UF, CEP 12345-678",
        "{{VALOR_INSCRITO}}": "R$ 100.000,00",
        "{{HONORARIOS}}": "R$ 30.000,00",
        "{{LIQUIDO_CLIENTE}}": "R$ 70.000,00",
        "{{PROPOSTA_70}}": "R$ 49.000,00",
        # Adicionar outras variáveis conforme o modelo
    }
    template_file = "/home/ubuntu/upload/Contrato_OVERTIME_Modelo.docx" # Caminho do modelo
    output_file = "/home/ubuntu/contrato_gerado_teste.docx"
    
    # Para salvar em arquivo
    # success = generate_contract_docx(template_file, output_file, mock_context)
    # if success:
    #     print(f"Contrato gerado com sucesso em: {output_file}")
    # else:
    #     print("Falha ao gerar contrato.")

    # Para salvar em buffer (útil para download em Flask)
    buffer = io.BytesIO()
    success_buffer = generate_contract_docx(template_file, buffer, mock_context)
    if success_buffer:
        print("Contrato gerado com sucesso em buffer.")
        # O buffer.getvalue() contém os bytes do arquivo .docx
        # Pode ser usado com send_file no Flask
        # with open("contrato_buffer_teste.docx", "wb") as f:
        #     f.write(buffer.getvalue())
        # print("Contrato do buffer salvo para teste.")
    else:
        print("Falha ao gerar contrato em buffer.")


